"""Document redaction engine.

Applies permanent redactions to PDF, DOCX, and XLSX files.  Detection runs
per text-unit (page / paragraph / cell) so that word boundaries from the
detector patterns are preserved, and document metadata (author, etc.) is
scrubbed as part of every redaction.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

from .detector import CATEGORY_MAP, Match, detect

# ── optional dependencies ──────────────────────────────────────────────────

try:
    import pymupdf as fitz  # PyMuPDF >= 1.26 preferred name
    _PDF_OK = True
except ImportError:
    try:
        import fitz  # legacy PyMuPDF import name
        _PDF_OK = True
    except Exception:
        _PDF_OK = False

try:
    import docx
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False

try:
    import openpyxl
    _XLSX_OK = True
except Exception:
    _XLSX_OK = False

SUPPORTED_EXTENSIONS: set[str] = {".pdf", ".docx", ".xlsx"}

DEFAULT_REPLACEMENT = "[REDACTED]"


@dataclass
class RedactResult:
    """Outcome of redacting one file."""
    src: Path
    outputs: list[Path] = field(default_factory=list)
    redaction_count: int = 0
    per_category: dict[str, int] = field(default_factory=dict)
    error: str | None = None


@dataclass
class ScanOptions:
    """Detection parameters shared by scan and redact operations."""
    enabled_categories: list[str] | None = None
    custom_patterns: list[str] | None = None
    custom_texts: list[str] | None = None
    replacement: str = DEFAULT_REPLACEMENT

    def run_detection(self, text: str) -> list[Match]:
        return detect(text, self.enabled_categories,
                      self.custom_patterns, self.custom_texts)


# ═══════════════════════════════════════════════════════════════════════════
# PDF
# ═══════════════════════════════════════════════════════════════════════════

def _redact_pdf(src: Path, out: Path, opts: ScanOptions) -> tuple[int, dict[str, int]]:
    """Permanently redact a PDF via PyMuPDF redaction annotations."""
    doc = fitz.open(str(src))
    total = 0
    per_cat: dict[str, int] = {}

    for page in doc:
        page_text = page.get_text()
        if not page_text.strip():
            continue
        matches = opts.run_detection(page_text)
        # Unique needles for this page, longest first (a longer match
        # subsumes shorter ones that start inside it).
        needles = sorted({m.text for m in matches}, key=len, reverse=True)
        for needle in needles:
            rects = page.search_for(needle)
            if not rects and " " in needle.strip():
                # Long needles that wrap across lines may not be found
                # whole — fall back to word-triple chunks.
                words = needle.split()
                for i in range(len(words) - 2):
                    rects.extend(page.search_for(" ".join(words[i:i + 3])))
            for rect in rects:
                page.add_redact_annot(rect, fill=(0, 0, 0))
                total += 1
        for m in matches:
            per_cat[m.category] = per_cat.get(m.category, 0) + 1
        # PDF_REDACT_IMAGE_NONE keeps embedded images intact; text under the
        # redaction rect is removed from the content stream.
        page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_NONE)

    # Scrub metadata (author, creator, etc. can leak identities).
    try:
        doc.set_metadata({})
        doc.del_xml_metadata()
    except Exception:
        pass

    doc.save(str(out), garbage=4, deflate=True)
    doc.close()
    return total, per_cat


# ═══════════════════════════════════════════════════════════════════════════
# DOCX
# ═══════════════════════════════════════════════════════════════════════════

def _mask_paragraph(paragraph, opts: ScanOptions) -> tuple[int, dict[str, int]]:
    """Redact matches inside one paragraph, preserving run formatting.

    Match coordinates are computed against the original paragraph text, then
    applied to each run in a single pass, so offsets never shift mid-edit.
    Returns (count, per-category counts).
    """
    text = paragraph.text
    if not text.strip():
        return 0, {}
    matches = opts.run_detection(text)
    if not matches:
        return 0, {}

    # Collect runs, including those inside hyperlinks when available.
    runs = list(paragraph.runs)
    for hl in getattr(paragraph, "hyperlinks", []):
        runs.extend(hl.runs)
    if not runs:
        return 0, {}

    # Map paragraph coordinates → run coordinates.
    spans: list[tuple[int, int, object]] = []  # (start, end, run)
    pos = 0
    for r in runs:
        spans.append((pos, pos + len(r.text), r))
        pos += len(r.text)

    # Gather per-run edit intervals in original coordinates.  A match that
    # spans runs inserts the replacement token in its first run and blanks
    # the continuation runs.
    run_edits: dict[int, list[tuple[int, int, str]]] = {}
    per_cat: dict[str, int] = {}
    masked = 0
    for m in matches:
        inserted = False
        for idx, (start, end, _run) in enumerate(spans):
            if end <= m.start or start >= m.end:
                continue
            ls = max(m.start, start) - start
            le = min(m.end, end) - start
            repl = opts.replacement if not inserted else ""
            run_edits.setdefault(idx, []).append((ls, le, repl))
            inserted = True
        if inserted:
            per_cat[m.category] = per_cat.get(m.category, 0) + 1
            masked += 1

    for idx, edits in run_edits.items():
        run = spans[idx][2]
        txt = run.text
        edits.sort()
        # Merge overlapping intervals (keep the first replacement token).
        merged: list[list] = []
        for ls, le, repl in edits:
            if merged and ls <= merged[-1][1]:
                merged[-1][1] = max(merged[-1][1], le)
            else:
                merged.append([ls, le, repl])
        new_parts: list[str] = []
        cursor = 0
        for ls, le, repl in merged:
            new_parts.append(txt[cursor:ls])
            new_parts.append(repl)
            cursor = le
        new_parts.append(txt[cursor:])
        run.text = "".join(new_parts)
    return masked, per_cat


def _redact_docx(src: Path, out: Path, opts: ScanOptions) -> tuple[int, dict[str, int]]:
    document = docx.Document(str(src))
    total = 0
    per_cat: dict[str, int] = {}

    def _accumulate(count: int, cats: dict[str, int]) -> None:
        nonlocal total
        total += count
        for k, v in cats.items():
            per_cat[k] = per_cat.get(k, 0) + v

    def _process_paragraphs(paragraphs) -> None:
        for para in paragraphs:
            _accumulate(*_mask_paragraph(para, opts))

    def _process_tables(tables) -> None:
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    _process_paragraphs(cell.paragraphs)
                    _process_tables(cell.tables)  # nested tables

    # Body
    _process_paragraphs(document.paragraphs)
    _process_tables(document.tables)
    # Headers / footers (often contain names, case numbers, dates)
    for section in document.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            if part is None:
                continue
            _process_paragraphs(part.paragraphs)
            _process_tables(part.tables)

    # Scrub document properties.
    try:
        cp = document.core_properties
        for attr in ("author", "last_modified_by", "title", "subject",
                     "keywords", "comments", "category"):
            try:
                setattr(cp, attr, "")
            except Exception:
                pass
    except Exception:
        pass

    document.save(str(out))
    return total, per_cat


# ═══════════════════════════════════════════════════════════════════════════
# XLSX
# ═══════════════════════════════════════════════════════════════════════════

# Header labels → detection category.  Spreadsheets like HR exports store
# PII in dedicated columns ("First Name", "Empl ID", numeric IDs, real
# dates) where per-cell text detection cannot help — so when a category is
# enabled, every value under a matching header is redacted wholesale.
_HEADER_MAP: list[tuple[re.Pattern, str]] = [
    (re.compile(r"^\s*(?:(?:first|last|full|middle|preferred|maiden|nick\s?|"
                r"supervisor|employee|student|patient|client|parent|guardian|"
                r"emergency\s+contact)\s*name|name)\s*$", re.IGNORECASE), "names"),
    (re.compile(r"e-?mail", re.IGNORECASE), "emails"),
    (re.compile(r"^\s*(?:phone|mobile|cell|telephone|fax)", re.IGNORECASE), "phones"),
    (re.compile(r"ssn|social\s+security", re.IGNORECASE), "ssn"),
    (re.compile(r"^\s*(?:dob|date\s+of\s+birth|birth\s*date)\s*$", re.IGNORECASE), "dates"),
    (re.compile(r"^\s*(?:street\s*)?address\s*\d*$", re.IGNORECASE), "addresses"),
    (re.compile(r"^\s*(?:zip|zip\s*code|postal\s*code)\s*$", re.IGNORECASE), "addresses"),
    (re.compile(r"^\s*(?:(?:empl|employee|student|patient|client|case|mrn)\s*"
                r"(?:id|#|no\.?|number)?|id)\s*$", re.IGNORECASE), "unique_ids"),
]

_HEADER_SCAN_ROWS = 10


def _header_columns(ws, enabled: set[str]) -> dict[int, tuple[int, str]]:
    """Return ``{column_index: (header_row, category)}`` for PII columns."""
    cols: dict[int, tuple[int, str]] = {}
    for row in ws.iter_rows(min_row=1, max_row=_HEADER_SCAN_ROWS):
        for cell in row:
            value = cell.value
            if not isinstance(value, str):
                continue
            text = value.strip()
            if not text or len(text) > 60:
                continue
            for pat, cat in _HEADER_MAP:
                if cat in enabled and pat.search(text):
                    cols.setdefault(cell.column, (cell.row, cat))
                    break
    return cols


def _process_xlsx(src: Path, opts: ScanOptions,
                  out: Path | None = None) -> tuple[int, dict[str, int]]:
    """Scan (out=None) or redact (out=path) a workbook with one procedure,
    keeping scan counts and redaction counts identical."""
    wb = openpyxl.load_workbook(str(src))
    total = 0
    per_cat: dict[str, int] = {}
    enabled = (set(opts.enabled_categories) if opts.enabled_categories
               else set(CATEGORY_MAP))

    for ws in wb.worksheets:
        header_cols = _header_columns(ws, enabled)
        for row in ws.iter_rows():
            for cell in row:
                # Header-mapped column: redact the whole value (any type).
                hc = header_cols.get(cell.column)
                if hc and cell.row > hc[0] and cell.value is not None \
                        and str(cell.value).strip():
                    per_cat[hc[1]] = per_cat.get(hc[1], 0) + 1
                    total += 1
                    if out is not None:
                        cell.value = opts.replacement
                    continue
                # Otherwise: per-cell text detection.
                value = cell.value
                if not isinstance(value, str) or not value.strip():
                    continue
                matches = opts.run_detection(value)
                if not matches:
                    continue
                new_value = value
                for m in matches:
                    per_cat[m.category] = per_cat.get(m.category, 0) + 1
                    total += 1
                    if out is not None:
                        pattern = re.compile(rf"\b{re.escape(m.text)}\b",
                                             re.IGNORECASE)
                        new_value = pattern.sub(opts.replacement, new_value)
                if out is not None:
                    cell.value = new_value

    if out is not None:
        # Scrub workbook properties.
        try:
            props = wb.properties
            props.creator = ""
            props.lastModifiedBy = ""
            props.title = ""
            props.subject = ""
            props.description = ""
        except Exception:
            pass
        wb.save(str(out))
    wb.close()
    return total, per_cat


def _redact_xlsx(src: Path, out: Path, opts: ScanOptions) -> tuple[int, dict[str, int]]:
    return _process_xlsx(src, opts, out=out)


# ═══════════════════════════════════════════════════════════════════════════
# Image output
# ═══════════════════════════════════════════════════════════════════════════

def _pdf_to_images(pdf_path: Path, out_dir: Path, fmt: str = "PNG") -> list[Path]:
    """Render PDF pages to images using PyMuPDF."""
    doc = fitz.open(str(pdf_path))
    out_dir.mkdir(parents=True, exist_ok=True)
    stem = pdf_path.stem
    paths: list[Path] = []
    for i, page in enumerate(doc, 1):
        pix = page.get_pixmap(dpi=150)
        if fmt.upper() in ("JPG", "JPEG"):
            out = out_dir / f"{stem}_p{i:03d}.jpg"
            pix.save(str(out), output="jpeg", jpg_quality=90)
        else:
            out = out_dir / f"{stem}_p{i:03d}.png"
            pix.save(str(out))
        paths.append(out)
    doc.close()
    return paths


# ═══════════════════════════════════════════════════════════════════════════
# Public API
# ═══════════════════════════════════════════════════════════════════════════

def extract_text(src_path: str | Path) -> str:
    """Extract full text from a document (used for scan previews)."""
    src = Path(src_path)
    ext = src.suffix.lower()

    if ext == ".pdf":
        if not _PDF_OK:
            raise RuntimeError("PyMuPDF (fitz) is required for PDF files.")
        doc = fitz.open(str(src))
        text = "\n".join(page.get_text() for page in doc)
        doc.close()
        return text

    if ext == ".docx":
        if not _DOCX_OK:
            raise RuntimeError("python-docx is required for DOCX files.")
        document = docx.Document(str(src))
        parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]

        def _walk_tables(tables) -> None:
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text)
                        _walk_tables(cell.tables)

        _walk_tables(document.tables)
        for section in document.sections:
            for part in (section.header, section.footer):
                if part is not None:
                    parts.extend(p.text for p in part.paragraphs if p.text.strip())
        return "\n".join(parts)

    if ext == ".xlsx":
        if not _XLSX_OK:
            raise RuntimeError("openpyxl is required for XLSX files.")
        wb = openpyxl.load_workbook(str(src), data_only=True)
        parts = [
            str(cell.value)
            for ws in wb.worksheets
            for row in ws.iter_rows()
            for cell in row
            if cell.value is not None and str(cell.value).strip()
        ]
        wb.close()
        return "\n".join(parts)

    raise ValueError(f"Unsupported file type: {ext}")


def redact_file(src_path: str | Path,
                opts: ScanOptions,
                out_path: str | Path | None = None,
                out_dir: str | Path | None = None,
                suffix: str = "_REDACTED",
                image_output: bool = False,
                image_format: str = "PNG") -> RedactResult:
    """Redact one document and save the result.

    Writes ``<stem><suffix><ext>`` next to the source by default, or into
    ``out_dir``.  When ``image_output`` is True and the source is a PDF, the
    redacted PDF is additionally rendered to per-page images.
    """
    src = Path(src_path)
    ext = src.suffix.lower()
    result = RedactResult(src=src)

    try:
        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")
        missing = {".pdf": (not _PDF_OK, "PyMuPDF (fitz)"),
                   ".docx": (not _DOCX_OK, "python-docx"),
                   ".xlsx": (not _XLSX_OK, "openpyxl")}[ext]
        if missing[0]:
            raise RuntimeError(f"{missing[1]} is not installed — required for {ext} files.")

        if out_path is None:
            if out_dir:
                out_path = Path(out_dir) / f"{src.stem}{suffix}{ext}"
            else:
                out_path = src.with_name(f"{src.stem}{suffix}{ext}")
        out_path = Path(out_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)

        if src.resolve() == out_path.resolve():
            raise ValueError("Output path equals the source file — refusing to overwrite.")

        if ext == ".pdf":
            count, per_cat = _redact_pdf(src, out_path, opts)
        elif ext == ".docx":
            count, per_cat = _redact_docx(src, out_path, opts)
        else:
            count, per_cat = _redact_xlsx(src, out_path, opts)

        result.redaction_count = count
        result.per_category = per_cat
        result.outputs.append(out_path)

        if image_output and ext == ".pdf":
            img_dir = (Path(out_dir) / f"{out_path.stem}_images") if out_dir \
                else out_path.parent / f"{out_path.stem}_images"
            result.outputs.extend(_pdf_to_images(out_path, img_dir, image_format))

    except Exception as exc:  # noqa: BLE001 — surface per-file errors
        result.error = str(exc)
    return result


def iter_text_units(src_path: str | Path) -> list[str]:
    """Extract text in the same units the redactor processes (per page /
    paragraph / cell), so scan counts match redaction counts exactly."""
    src = Path(src_path)
    ext = src.suffix.lower()

    if ext == ".pdf":
        if not _PDF_OK:
            raise RuntimeError("PyMuPDF (fitz) is required for PDF files.")
        doc = fitz.open(str(src))
        units = [page.get_text() for page in doc]
        doc.close()
        return units

    if ext == ".docx":
        if not _DOCX_OK:
            raise RuntimeError("python-docx is required for DOCX files.")
        document = docx.Document(str(src))
        units: list[str] = []

        def _walk_tables(tables) -> None:
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        units.extend(p.text for p in cell.paragraphs)
                        _walk_tables(cell.tables)

        units.extend(p.text for p in document.paragraphs)
        _walk_tables(document.tables)
        for section in document.sections:
            for part in (section.header, section.footer):
                if part is not None:
                    units.extend(p.text for p in part.paragraphs)
        return units

    if ext == ".xlsx":
        if not _XLSX_OK:
            raise RuntimeError("openpyxl is required for XLSX files.")
        wb = openpyxl.load_workbook(str(src), data_only=True)
        units = [
            str(cell.value)
            for ws in wb.worksheets
            for row in ws.iter_rows()
            for cell in row
            if isinstance(cell.value, str) and cell.value.strip()
        ]
        wb.close()
        return units

    raise ValueError(f"Unsupported file type: {ext}")


def scan_file(src_path: str | Path, opts: ScanOptions) -> dict[str, int]:
    """Return per-category detection counts for one file.

    Uses the same per-unit procedure as redaction so the previewed counts
    match what a redaction pass will actually mask.
    """
    src = Path(src_path)
    if src.suffix.lower() == ".xlsx":
        _total, per_cat = _process_xlsx(src, opts, out=None)
        return dict(sorted(per_cat.items()))
    counts: dict[str, int] = {}
    for unit in iter_text_units(src_path):
        for m in opts.run_detection(unit):
            counts[m.category] = counts.get(m.category, 0) + 1
    return dict(sorted(counts.items()))

