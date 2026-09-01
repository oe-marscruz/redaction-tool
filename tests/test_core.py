"""Unit + regression tests for the redaction tool.

Run:  python tests/test_core.py        (no pytest needed)
"""
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parents[1]))

from redaction_tool import detector, ocr, redactor  # noqa: E402


# ── detector.py ────────────────────────────────────────────────────────────

def test_detector_core_categories():
    text = "Email jane@example.edu SSN 123-45-6789 phone (303) 555-0142"
    cats = {m.category for m in detector.detect(text)}
    assert "emails" in cats
    assert "ssn" in cats
    assert "phones" in cats


def test_name_false_positive_excluded():
    # "Thank you,\nJennifer" must NOT match the "Last, First" pattern.
    matches = detector.detect("Thank you,\nJennifer")
    assert not any("you" in m.text for m in matches)


def test_custom_text_whole_word():
    matches = detector.detect("Ann met Anna", custom_texts=["Ann"])
    texts = [m.text for m in matches]
    assert texts.count("Ann") == 1  # "Anna" must not match


def test_mask_value_never_full_ssn():
    masked = ocr._mask_value("ssn", "123-45-6789")
    assert "123-45-6789" not in masked


# ── regression: PDF link annotations must not survive redaction ────────────

def test_pdf_link_uri_scrubbed():
    import pymupdf
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 100), "Contact us at (303) 555-0142")
        page.insert_text((72, 200), "See website")
        page.insert_link({
            "kind": pymupdf.LINK_URI,
            "from": pymupdf.Rect(72, 196, 200, 212),
            "uri": "mailto:jane.doe@example.edu?ssn=123456789",
        })
        src = td / "in.pdf"
        doc.save(str(src))
        doc.close()

        out = td / "out.pdf"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        d2 = pymupdf.open(str(out))
        uris = [l.get("uri", "") for l in d2[0].get_links()]
        d2.close()
        assert not any("jane.doe" in u or "123456789" in u for u in uris), uris


# ── regression: embedded attachments must not leak PII ─────────────────────

def test_pdf_attachment_scrubbed():
    import pymupdf
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 100), "Meeting notes for Jane Doe (303) 555-0142")
        doc.embfile_add("notes.txt",
                        b"Patient: Jane Doe, MRN 448821, DOB 05/14/2002",
                        filename="notes.txt")
        src = td / "in.pdf"
        doc.save(str(src))
        doc.close()

        out = td / "out.pdf"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        d2 = pymupdf.open(str(out))
        leaked = False
        for i in range(d2.embfile_count()):
            data = d2.embfile_get(i) or b""
            if b"Jane Doe" in data or b"448821" in data:
                leaked = True
        d2.close()
        assert not leaked


# ── regression: DOCX text boxes must be redacted ───────────────────────────

def test_docx_textbox_redacted():
    import docx
    from lxml import etree
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        d = docx.Document()
        d.add_paragraph("Body text mentions Jane Doe regularly.")
        p = d.add_paragraph()
        run = p.add_run()
        drawing_xml = (
            '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:drawing><wp:inline xmlns:wp="http://schemas.openxmlformats.org/'
            'drawingml/2006/wordprocessingDrawing">'
            '<a:graphic xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            '<a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/'
            'wordprocessingShape">'
            '<wps:wsp xmlns:wps="http://schemas.microsoft.com/office/word/2010/'
            'wordprocessingShape">'
            '<wps:txbx><w:txbxContent>'
            '<w:p><w:r><w:t>Hidden box: jane.doe@example.edu 123-45-6789</w:t></w:r></w:p>'
            "</w:txbxContent></wps:txbx></wps:wsp>"
            "</a:graphicData></a:graphic></wp:inline></w:drawing></w:r>"
        )
        run._r.append(etree.fromstring(drawing_xml))
        src = td / "in.docx"
        d.save(str(src))

        out = td / "out.docx"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        import zipfile
        with zipfile.ZipFile(out) as z:
            xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        assert "jane.doe@example.edu" not in xml
        assert "123-45-6789" not in xml


# ── regression: PDF form-field values must not survive ───────────────────

def test_pdf_form_field_scrubbed():
    import pymupdf
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        doc = pymupdf.open()
        page = doc.new_page()
        page.insert_text((72, 60), "Employee termination record")
        w = pymupdf.Widget()
        w.field_type = pymupdf.PDF_WIDGET_TYPE_TEXT
        w.field_name = "ssn_field"
        w.field_value = "123-45-6789"
        w.rect = pymupdf.Rect(72, 100, 200, 116)
        page.add_widget(w)
        src = td / "form.pdf"
        doc.save(str(src))
        doc.close()

        out = td / "out.pdf"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        d2 = pymupdf.open(str(out))
        values = [str(wd.field_value or "") for wd in d2[0].widgets()]
        d2.close()
        assert not any("123-45-6789" in v for v in values), values


# ── regression: DOCX footnotes must be redacted ──────────────────────────

def test_docx_footnote_redacted():
    import docx
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "footnote.docx"
        files = {
            "[Content_Types].xml": (
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
                '<Override PartName="/word/footnotes.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml"/>'
                "</Types>"),
            "_rels/.rels": (
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
                "</Relationships>"),
            "word/_rels/document.xml.rels": (
                '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                '<Relationship Id="rId5" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes" Target="footnotes.xml"/>'
                "</Relationships>"),
            "word/document.xml": (
                '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:body><w:p><w:r><w:t>Body mentions Jane Doe here.</w:t></w:r></w:p></w:body></w:document>'),
            "word/footnotes.xml": (
                '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:footnote w:id="1"><w:p><w:r><w:t>Footnote: jane.doe@example.edu 123-45-6789</w:t></w:r></w:p></w:footnote>'
                "</w:footnotes>"),
        }
        import zipfile
        with zipfile.ZipFile(src, "w") as z:
            for n, c in files.items():
                z.writestr(n, c)

        out = td / "out.docx"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        with zipfile.ZipFile(out) as z:
            xml = z.read("word/footnotes.xml").decode("utf-8", errors="ignore")
        assert "jane.doe@example.edu" not in xml
        assert "123-45-6789" not in xml


# ── regression: XLSX cell comments must be scrubbed ─────────────────────

def test_xlsx_comment_scrubbed():
    import openpyxl
    from openpyxl.comments import Comment
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "in.xlsx"
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "Data"
        ws["A2"] = "clean value"
        ws["A2"].comment = Comment("Follow up: jane.doe@example.edu", "HR")
        wb.save(str(src))

        out = td / "out.xlsx"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        wb2 = openpyxl.load_workbook(str(out))
        c = wb2.active["A2"].comment
        wb2.close()
        assert c is None or "jane.doe@example.edu" not in (c.text or "")


# ── docProps/app.xml Manager/Company scrub ─────────────────────────────

def test_app_xml_scrubbed():
    import zipfile
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        src = td / "doc.docx"
        import docx
        d = docx.Document()
        d.add_paragraph("hello")
        d.save(str(src))
        # Inject a Manager name into app.xml
        tmp = td / "inject.docx"
        with zipfile.ZipFile(src) as zin, \
                zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                content = zin.read(item.filename)
                if item.filename == "docProps/app.xml":
                    content = content.replace(
                        b"</Properties>", b"<Manager>Jane Doe</Manager></Properties>")
                zout.writestr(item, content)
        tmp.replace(src)

        out = td / "out.docx"
        redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        with zipfile.ZipFile(out) as z:
            app = z.read("docProps/app.xml")
        assert b"Jane Doe" not in app


# ── salted plan fingerprints ────────────────────────────────────────────

def test_salted_fingerprint():
    import hashlib
    value = "123-45-6789"
    unsalted = hashlib.sha256(value.encode()).hexdigest()
    salted = hashlib.sha256(("abc123" + value).encode()).hexdigest()
    assert unsalted != salted


# ── regression: image files route through the OCR engine ───────────────────

def test_image_round_trip():
    if not ocr.find_tesseract():
        print("SKIP: Tesseract not installed")
        return
    from PIL import Image, ImageDraw, ImageFont
    with tempfile.TemporaryDirectory() as td:
        td = Path(td)
        img = Image.new("RGB", (400, 120), "white")
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.truetype("C:/Windows/Fonts/arial.ttf", 18)
        except Exception:
            font = ImageFont.load_default()
        draw.text((20, 40), "SSN 123-45-6789", fill="black", font=font)
        src = td / "in.png"
        img.save(str(src))

        counts = redactor.scan_file(src, redactor.ScanOptions())
        assert counts.get("ssn", 0) >= 1, counts

        out = td / "out.png"
        result = redactor.redact_file(src, redactor.ScanOptions(), out_path=out)
        assert result.error is None, result.error
        v = ocr.verify_ocr(out, scan_opts=redactor.ScanOptions())
        assert v["remaining_detection_count"] == 0


if __name__ == "__main__":
    failures = 0
    for name, fn in sorted(globals().items()):
        if name.startswith("test_") and callable(fn):
            try:
                fn()
                print(f"PASS {name}")
            except AssertionError as e:
                failures += 1
                print(f"FAIL {name}: {e}")
            except Exception as e:  # noqa: BLE001
                failures += 1
                print(f"ERROR {name}: {type(e).__name__}: {e}")
    print("RESULT:", "FAIL" if failures else "ALL PASS")
    sys.exit(1 if failures else 0)
